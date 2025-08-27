"""Meeting minutes pipeline implementing transcription, diarization, summarization and export utilities."""

from __future__ import annotations

import argparse
import json
import logging
import os
import re
import subprocess
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional

from pydantic import BaseModel, ValidationError


class Decision(BaseModel):
    text: str
    rationale: str


class ActionItem(BaseModel):
    owner: str
    task: str
    due: str


class Mention(BaseModel):
    entity: str
    context: str


class ChunkSummary(BaseModel):
    chunk_window: Dict[str, str]
    topic_title: str
    key_points: List[str]
    decisions: List[Decision]
    action_items: List[ActionItem]
    risks_or_blockers: List[str]
    open_questions: List[str]
    mentions: List[Mention]


# ---------------------------------------------------------------------------
# Audio/Transcription utilities
# ---------------------------------------------------------------------------

def extract_audio(video_path: str, wav_path: str) -> None:
    """Extract mono 16 kHz PCM audio from a video file using ffmpeg.

    Args:
        video_path: Path to the input video file.
        wav_path: Destination path for the extracted WAV file.

    Raises:
        RuntimeError: If ffmpeg is not installed or extraction fails.
    """

    cmd = [
        "ffmpeg",
        "-y",
        "-i",
        video_path,
        "-ac",
        "1",
        "-ar",
        "16000",
        "-f",
        "wav",
        wav_path,
    ]
    logging.info("Extracting audio with ffmpeg")
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except FileNotFoundError as exc:
        raise RuntimeError("ffmpeg is required but was not found") from exc
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(f"ffmpeg failed: {exc.stderr.decode(errors='ignore')}") from exc


def transcribe_and_diarize(
    wav_path: str,
    language: Optional[str],
    max_speakers: int,
    asr_model: str = "medium",
    compute_type: str = "int8",
    device: str = "cpu",
    cpu_threads: Optional[int] = None,
    skip_align: bool = False,
    no_diarization: bool = False,
    hf_token: Optional[str] = None,
) -> Dict[str, Any]:
    """Transcribe and optionally diarize an audio file using WhisperX.

    Args:
        wav_path: Path to the WAV audio file.
        language: Optional language code to force transcription language.
        max_speakers: Maximum number of speakers for diarization.
        asr_model: Whisper model size.
        compute_type: faster-whisper compute precision.
        device: Target device ("cpu" or "cuda" or "auto").
        cpu_threads: Number of CPU threads to use for ASR.
        skip_align: If True, skip word-level alignment for speed.
        no_diarization: If True, skip speaker diarization.
        hf_token: Optional HuggingFace token for diarization.

    Returns:
        Dictionary with a ``segments`` list describing speaker segments.
    """

    import torch  # type: ignore
    import whisperx  # type: ignore

    if device == "auto":
        device = "cuda" if torch.cuda.is_available() else "cpu"


    logging.info(
        "Loading WhisperX model %s on %s (%s)", asr_model, device, compute_type

    logging.info("Loading WhisperX model %s on %s (%s)", asr_model, device, compute_type)
    model = whisperx.load_model(
        asr_model, device=device, compute_type=compute_type, cpu_threads=cpu_threads

    )
    try:
        model = whisperx.load_model(
            asr_model, device=device, compute_type=compute_type, cpu_threads=cpu_threads
        )
    except TypeError:
        logging.warning(
            "Installed whisperx does not support 'cpu_threads'; loading without it"
        )
        model = whisperx.load_model(
            asr_model, device=device, compute_type=compute_type
        )

    asr_kwargs = {}
    if language:
        asr_kwargs["language"] = language
    asr_kwargs.update(dict(vad_filter=True))
    result = model.transcribe(wav_path, **asr_kwargs)

    audio = whisperx.load_audio(wav_path)

    if not skip_align:
        model_a, metadata = whisperx.load_align_model(
            language_code=result["language"], device=device
        )
        aligned = whisperx.align(
            result["segments"], model_a, metadata, audio, device=device
        )
    else:
        aligned = {"segments": result["segments"]}

    if not no_diarization:
        diarize_model = whisperx.diarize.DiarizationPipeline(
            use_auth_token=hf_token, device=device
        )
        diarize_segments = diarize_model(audio, max_speakers=max_speakers)
        aligned = whisperx.assign_word_speakers(diarize_segments, aligned)

    segments = []
    for seg in aligned["segments"]:
        segments.append(
            {
                "start": float(seg["start"]),
                "end": float(seg["end"]),
                "speaker": seg.get("speaker", "SPEAKER_00"),
                "text": seg["text"].strip(),
            }
        )
    return {"segments": segments}


# ---------------------------------------------------------------------------
# Segments manipulation
# ---------------------------------------------------------------------------

def stitch_segments(
    segments: List[Dict[str, Any]], min_secs: float = 0.5, max_chars: int = 600
) -> List[Dict[str, Any]]:
    """Merge tiny fragments and stitch consecutive same-speaker segments.

    Args:
        segments: Raw diarized segments.
        min_secs: Minimum segment duration in seconds; shorter segments will be
            merged into the previous segment.
        max_chars: Maximum character length for stitched speaker blocks.

    Returns:
        List of processed segments.
    """

    if not segments:
        return []

    merged: List[Dict[str, Any]] = []
    for seg in segments:
        dur = seg["end"] - seg["start"]
        if merged and dur < min_secs:
            prev = merged[-1]
            prev["end"] = seg["end"]
            prev["text"] = (prev["text"] + " " + seg["text"]).strip()
        else:
            merged.append(dict(seg))

    stitched: List[Dict[str, Any]] = []
    for seg in merged:
        if (
            stitched
            and seg["speaker"] == stitched[-1]["speaker"]
            and len(stitched[-1]["text"]) + len(seg["text"]) <= max_chars
        ):
            prev = stitched[-1]
            prev["end"] = seg["end"]
            prev["text"] = (prev["text"] + " " + seg["text"]).strip()
        else:
            stitched.append(dict(seg))
    return stitched


def make_chunks(stitched: List[Dict[str, Any]], chunk_seconds: int) -> List[Dict[str, Any]]:
    """Slice stitched segments into fixed-size time chunks.

    Args:
        stitched: Stitched speaker segments.
        chunk_seconds: Chunk duration in seconds.

    Returns:
        List of chunk descriptors each containing entries that overlap the chunk.
    """

    if not stitched:
        return []

    total_end = max(seg["end"] for seg in stitched)
    chunks: List[Dict[str, Any]] = []
    idx = 0
    start = 0.0
    while start < total_end:
        end = start + chunk_seconds
        entries: List[Dict[str, Any]] = []
        for seg in stitched:
            if seg["end"] <= start or seg["start"] >= end:
                continue
            entries.append(dict(seg))
        chunks.append({"idx": idx, "start": start, "end": min(end, total_end), "entries": entries})
        start = end
        idx += 1
    return chunks


def ts(sec: float) -> str:
    """Format seconds as ``HH:MM:SS`` string."""
    sec_int = int(sec)
    h, rem = divmod(sec_int, 3600)
    m, s = divmod(rem, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


# ---------------------------------------------------------------------------
# Summarization and merge
# ---------------------------------------------------------------------------

def summarize_chunk(
    chunk: Dict[str, Any],
    attendees: Optional[List[str]],
    system_prompt: str,
    model: str,
    api_key: str,
) -> Dict[str, Any]:
    """Summarize a diarized chunk using an LLM with strict JSON schema.

    Args:
        chunk: Chunk descriptor from :func:`make_chunks`.
        attendees: Optional list of known attendees.
        system_prompt: System prompt text.
        model: Model name for OpenAI completion.
        api_key: OpenAI API key.

    Returns:
        Parsed JSON summary validated by Pydantic.
    """

    from openai import OpenAI  # type: ignore

    start_ts = ts(chunk["start"])
    end_ts = ts(chunk["end"])
    participants = ",".join(attendees) if attendees else "unknown"

    lines = []
    for entry in chunk["entries"]:
        lines.append(
            f"{entry['speaker']} [{ts(entry['start'])} - {ts(entry['end'])}]: {entry['text']}"
        )
    transcript = "\n".join(lines)

    user_content = (
        f"Participants (if known): {participants}\n"
        f"Time window: {start_ts} - {end_ts}\n"
        f"Diarized transcript:\n{transcript}"
    )

    client = OpenAI(api_key=api_key)

    def _call(prompt: str) -> Dict[str, Any]:
        resp = client.responses.create(
            model=model,
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
        )
        return json.loads(resp.output_text)

    try:
        parsed = ChunkSummary.model_validate(_call(user_content))
        return parsed.model_dump()
    except (json.JSONDecodeError, ValidationError) as err:
        correction = user_content + "\nThe previous output was invalid: " + str(err)
        parsed = ChunkSummary.model_validate(_call(correction))
        return parsed.model_dump()


def merge_partials(partials: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Merge chunk summaries into final minutes structure.

    Args:
        partials: List of validated chunk summaries.

    Returns:
        Dictionary representing the merged minutes document.
    """

    agenda: List[str] = []
    by_window: List[Dict[str, Any]] = []
    decision_map: Dict[str, Dict[str, str]] = {}
    action_map: Dict[tuple, Dict[str, str]] = {}
    risks: List[str] = []
    open_qs: List[str] = []

    for part in partials:
        cw = part["chunk_window"]
        agenda.append(part.get("topic_title", ""))
        by_window.append(
            {
                "start": cw["start"],
                "end": cw["end"],
                "bullets": part.get("key_points", [])[:2],
            }
        )

        for dec in part.get("decisions", []):
            norm = re.sub(r"\s+", " ", dec["text"]).strip().lower()
            if norm not in decision_map:
                decision_map[norm] = dec

        for act in part.get("action_items", []):
            key = (
                act["owner"].strip().lower(),
                re.sub(r"\s+", " ", act["task"]).strip().lower(),
            )
            existing = action_map.get(key)
            due = act.get("due", "TBD")
            item = {
                "owner": act["owner"],
                "task": act["task"],
                "due": due,
                "window_start": cw["start"],
                "window_end": cw["end"],
            }
            if existing:
                if existing["due"] == "TBD" and due != "TBD":
                    existing["due"] = due
                elif (
                    existing["due"] != "TBD" and due != "TBD" and due < existing["due"]
                ):
                    existing["due"] = due
                continue
            action_map[key] = item

        risks.extend(part.get("risks_or_blockers", []))
        open_qs.extend(part.get("open_questions", []))

    merged = {
        "date": datetime.utcnow().date().isoformat(),
        "participants": [],
        "agenda": [a for a in agenda if a],
        "by_window": by_window,
        "decisions": list(decision_map.values()),
        "action_items": list(action_map.values()),
        "risks": list(dict.fromkeys(risks)),
        "open_questions": list(dict.fromkeys(open_qs)),
        "speaker_aliases": [],
    }
    return merged


# ---------------------------------------------------------------------------
# Rendering & export
# ---------------------------------------------------------------------------

def render_markdown(merged: Dict[str, Any]) -> str:
    """Render merged minutes structure to Markdown."""

    lines = ["# Meeting Minutes", f"Date: {merged['date']}", ""]
    lines.append("## Participants")
    if merged["participants"]:
        for p in merged["participants"]:
            lines.append(f"- {p}")
    else:
        lines.append("- Unknown")
    lines.append("")

    if merged["agenda"]:
        lines.append("## Agenda")
        for a in merged["agenda"]:
            lines.append(f"- {a}")
        lines.append("")

    lines.append("## Discussion by Time Window")
    for win in merged["by_window"]:
        lines.append(f"### {win['start']} - {win['end']}")
        for b in win.get("bullets", []):
            lines.append(f"- {b}")
        lines.append("")

    if merged["decisions"]:
        lines.append("## Decisions")
        for d in merged["decisions"]:
            lines.append(f"- {d['text']} (rationale: {d['rationale']})")
        lines.append("")

    if merged["action_items"]:
        lines.append("## Action Items")
        lines.append("| Owner | Task | Due | Source Window |")
        lines.append("| --- | --- | --- | --- |")
        for a in merged["action_items"]:
            lines.append(
                f"| {a['owner']} | {a['task']} | {a['due']} | {a['window_start']}–{a['window_end']} |"
            )
        lines.append("")

    if merged["risks"]:
        lines.append("## Risks/Blockers")
        for r in merged["risks"]:
            lines.append(f"- {r}")
        lines.append("")

    if merged["open_questions"]:
        lines.append("## Parking Lot / Open Questions")
        for q in merged["open_questions"]:
            lines.append(f"- {q}")
        lines.append("")

    return "\n".join(lines)


def export_docx(merged: Dict[str, Any], path: str) -> None:
    """Export merged minutes to a DOCX file."""

    from docx import Document  # type: ignore

    doc = Document()
    doc.add_heading("Meeting Minutes", level=1)
    doc.add_paragraph(f"Date: {merged['date']}")

    doc.add_heading("Participants", level=2)
    if merged["participants"]:
        for p in merged["participants"]:
            doc.add_paragraph(p, style="List Bullet")
    else:
        doc.add_paragraph("Unknown", style="List Bullet")

    if merged["agenda"]:
        doc.add_heading("Agenda", level=2)
        for a in merged["agenda"]:
            doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Discussion by Time Window", level=2)
    for win in merged["by_window"]:
        doc.add_heading(f"{win['start']} - {win['end']}", level=3)
        for b in win.get("bullets", []):
            doc.add_paragraph(b, style="List Bullet")

    if merged["decisions"]:
        doc.add_heading("Decisions", level=2)
        for d in merged["decisions"]:
            doc.add_paragraph(f"{d['text']} (rationale: {d['rationale']})", style="List Bullet")

    if merged["action_items"]:
        doc.add_heading("Action Items", level=2)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Owner"
        hdr_cells[1].text = "Task"
        hdr_cells[2].text = "Due"
        hdr_cells[3].text = "Source Window"
        for a in merged["action_items"]:
            row_cells = table.add_row().cells
            row_cells[0].text = a["owner"]
            row_cells[1].text = a["task"]
            row_cells[2].text = a["due"]
            row_cells[3].text = f"{a['window_start']}–{a['window_end']}"

    if merged["risks"]:
        doc.add_heading("Risks/Blockers", level=2)
        for r in merged["risks"]:
            doc.add_paragraph(r, style="List Bullet")

    if merged["open_questions"]:
        doc.add_heading("Parking Lot / Open Questions", level=2)
        for q in merged["open_questions"]:
            doc.add_paragraph(q, style="List Bullet")

    doc.save(path)


def export_actions_csv(merged: Dict[str, Any], path: str) -> None:
    """Export action items to CSV."""

    import pandas as pd  # type: ignore

    df = pd.DataFrame(merged["action_items"]).loc[
        :, ["owner", "task", "due", "window_start", "window_end"]
    ]
    df.to_csv(path, index=False)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Meeting minutes pipeline")
    parser.add_argument("--video", required=True, help="Path to meeting video")
    parser.add_argument("--audio", help="Optional pre-extracted audio")
    parser.add_argument("--chunk-seconds", type=int, default=180)
    parser.add_argument("--max-speakers", type=int, default=6)
    parser.add_argument("--attendees", help="Comma-separated known attendees")
    parser.add_argument("--language", help="Force language for transcription")
    parser.add_argument("--openai-api-key", help="OpenAI API key")
    parser.add_argument(
        "--model", default="gpt-4o", help="Text model for summarization"
    )
    parser.add_argument("--output-dir", default="./out")
    parser.add_argument("--asr-model", default="medium", help="ASR model size: tiny/base/small/medium/large-v3")
    parser.add_argument("--compute-type", default="int8", help="faster-whisper compute type on CPU: int8/float32")
    parser.add_argument("--device", default="cpu", help="cpu or cuda")
    parser.add_argument("--cpu-threads", type=int, default=os.cpu_count(), help="CPU threads for ASR")
    parser.add_argument("--skip-align", action="store_true", help="Skip word-level alignment (faster)")
    parser.add_argument("--no-diarization", action="store_true", help="Skip speaker diarization (much faster)")
    parser.add_argument("--hf-token", type=str, default=os.getenv("HF_TOKEN"), help="HuggingFace token for pyannote")
    args = parser.parse_args()

    logging.basicConfig(level=logging.INFO, format="%(levelname)s:%(message)s")

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    wav_path = Path(args.audio) if args.audio else out_dir / "meeting.wav"
    if not args.audio:
        extract_audio(args.video, str(wav_path))

    diarized = transcribe_and_diarize(
        str(wav_path),
        args.language,
        args.max_speakers,
        asr_model=args.asr_model,
        compute_type=args.compute_type,
        device=args.device,
        cpu_threads=args.cpu_threads,
        skip_align=args.skip_align,
        no_diarization=args.no_diarization,
        hf_token=args.hf_token,
    )
    (out_dir / "transcript_diarized.json").write_text(
        json.dumps(diarized, indent=2)
    )

    stitched = stitch_segments(diarized["segments"])
    chunks = make_chunks(stitched, args.chunk_seconds)

    system_prompt = Path("prompts/system_chunk.txt").read_text()
    attendees = [a.strip() for a in args.attendees.split(",")] if args.attendees else None
    api_key = args.openai_api_key or os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OpenAI API key not provided")

    partials: List[Dict[str, Any]] = []
    chunk_dir = out_dir / "chunks"
    chunk_dir.mkdir(exist_ok=True)
    for chunk in chunks:
        if not chunk["entries"]:
            summary = {"chunk_window": {"start": ts(chunk["start"]), "end": ts(chunk["end"])}}
        else:
            summary = summarize_chunk(
                chunk, attendees, system_prompt, args.model, api_key
            )
        (chunk_dir / f"chunk_{chunk['idx']:03d}.json").write_text(
            json.dumps(summary, indent=2)
        )
        partials.append(summary)

    merged = merge_partials(partials)
    if attendees:
        merged["participants"] = attendees
    speakers = sorted({e["speaker"] for c in chunks for e in c["entries"]})
    merged["speaker_aliases"] = [
        {"label": s, "name": None} for s in speakers
    ]

    (out_dir / "minutes_merged.json").write_text(json.dumps(merged, indent=2))

    md = render_markdown(merged)
    (out_dir / "minutes.md").write_text(md)
    export_docx(merged, str(out_dir / "minutes.docx"))
    export_actions_csv(merged, str(out_dir / "actions.csv"))


if __name__ == "__main__":  # pragma: no cover
    main()

